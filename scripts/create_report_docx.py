from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Reporte_Tecnico_LOCATIONMIKE.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243447"
MUTED = "667085"
LIGHT = "F4F6F9"
BORDER = "C9D5E2"
WHITE = "FFFFFF"


def set_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_font(first, bold=True, color=INK)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest, color=INK)
    else:
        run = p.add_run(text)
        set_font(run, color=INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    set_font(p.add_run(text), color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    set_font(p.add_run(text), color=INK)
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(f"{label}: "), bold=True, color=DARK_BLUE)
    set_font(p.add_run(text), color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_payload(doc):
    payload = '''{
  "usuario": "Nombre completo",
  "dispositivo": {
    "nombre": "Samsung Galaxy S24",
    "version": "Android 15"
  },
  "ubicacion": {
    "lat": 32.5149,
    "lon": -117.0382,
    "altitud": 112.5,
    "precision": 4.2
  },
  "movimiento": {
    "velocidad": 1.4,
    "marca_tiempo": "2026-07-04T14:58:00.000Z"
  },
  "estado": {
    "bateria": 0.85
  }
}'''
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EEF3F8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for i, line in enumerate(payload.splitlines()):
        if i:
            p.add_run().add_break()
        set_font(p.add_run(line), size=8.5, color="183B56", name="Consolas")


def add_evidence_box(doc, number, title, description):
    doc.add_heading(f"Captura {number}. {title}", level=2)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(40)
    set_font(p.add_run("INSERTAR CAPTURA AQUÍ"), size=14, bold=True, color=MUTED)
    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    set_font(caption.add_run(f"Figura {number}. {description}"), size=9, italic=True, color=MUTED)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.333

for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

# Encabezado y pie discretos.
header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("MIKELOCATIONS  |  Reporte técnico"), size=9, color=MUTED)
add_page_number(section.footer.paragraphs[0])

# Portada editorial.
doc.add_paragraph().paragraph_format.space_after = Pt(72)
kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(kicker.add_run("RETO MÓVIL"), size=11, bold=True, color=BLUE)
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(12)
title.paragraph_format.space_after = Pt(8)
set_font(title.add_run("Telemetría de Ubicación\npor Distancia"), size=28, bold=True, color=DARK_BLUE)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(42)
set_font(subtitle.add_run("React Native + Expo SDK 54 + n8n"), size=15, color=MUTED)

meta = doc.add_table(rows=5, cols=2)
set_table_geometry(meta, [2700, 6660])
meta_data = [
    ("Institución", "[ESCRIBE EL NOMBRE DE LA INSTITUCIÓN]"),
    ("Asignatura", "[ESCRIBE LA ASIGNATURA]"),
    ("Alumno", "[ESCRIBE TU NOMBRE COMPLETO]"),
    ("Docente", "[ESCRIBE EL NOMBRE DEL DOCENTE]"),
    ("Fecha", "9 de julio de 2026"),
]
for row, (label, value) in zip(meta.rows, meta_data):
    set_cell_shading(row.cells[0], "E8EEF5")
    set_font(row.cells[0].paragraphs[0].add_run(label), bold=True, color=DARK_BLUE)
    set_font(row.cells[1].paragraphs[0].add_run(value), color=INK)

doc.add_page_break()

doc.add_heading("Resumen ejecutivo", level=1)
add_body(doc, "Este reporte documenta el desarrollo de MIKELOCATIONS, una aplicación móvil construida con React Native y Expo SDK 54. La solución escucha cambios del sensor de ubicación, registra desplazamientos aproximados de 10 metros y envía telemetría extendida a un webhook de n8n mediante peticiones HTTP POST.")
add_callout(doc, "Resultado", "La transmisión se relaciona con el movimiento real del dispositivo, evitando solicitudes periódicas cuando el usuario permanece inmóvil.")

doc.add_heading("Datos de entrega", level=1)
data = doc.add_table(rows=4, cols=2)
set_table_geometry(data, [2700, 6660])
delivery = [
    ("Repositorio", "https://github.com/Bladi1802/LOCATIONMIKE"),
    ("Tecnología", "React Native, Expo SDK 54, TypeScript y n8n"),
    ("Intervalo", "10 metros de desplazamiento aproximado"),
    ("Plataformas", "Android e iOS"),
]
for row, (label, value) in zip(data.rows, delivery):
    set_cell_shading(row.cells[0], LIGHT)
    set_font(row.cells[0].paragraphs[0].add_run(label), bold=True, color=DARK_BLUE)
    set_font(row.cells[1].paragraphs[0].add_run(value), color=INK)

doc.add_heading("Objetivos", level=1)
add_bullet(doc, "Capturar coordenadas de ubicación con precisión alta.")
add_bullet(doc, "Disparar telemetría después de aproximadamente 10 metros de desplazamiento.")
add_bullet(doc, "Incluir información de movimiento, dispositivo, sistema y batería.")
add_bullet(doc, "Enviar el payload a n8n sin incorporar dependencias pesadas.")
add_bullet(doc, "Mantener rastreo en segundo plano bajo los permisos de cada plataforma.")

doc.add_heading("Fundamento conceptual", level=1)
doc.add_heading("Polling tradicional", level=2)
add_body(doc, "En un esquema de polling, la aplicación consulta o transmite información cada cierto intervalo de tiempo. Aunque es sencillo, puede despertar repetidamente el GPS, el procesador y la radio de red incluso cuando el teléfono no se ha desplazado. Esto genera payloads repetidos, mayor tráfico y consumo innecesario de batería.")
doc.add_heading("Suscripción reactiva por distancia", level=2)
add_body(doc, "La solución utiliza una suscripción de ubicación configurada con distanceInterval: 10. El proveedor del sistema operativo observa los cambios de posición y entrega una actualización cuando estima que se alcanzó la distancia solicitada. Si el usuario permanece inmóvil, la app no necesita enviar peticiones periódicas al webhook.")
add_callout(doc, "Precisión", "Los 10 metros son una referencia solicitada al proveedor, no una frontera geométrica exacta. La frecuencia real depende de la señal GPS, el entorno, la plataforma y sus políticas energéticas.")

doc.add_heading("Arquitectura de la solución", level=1)
add_number(doc, "La app solicita permiso de ubicación en primer plano y obtiene la posición inicial.")
add_number(doc, "watchPositionAsync actualiza el mapa mientras la interfaz permanece abierta.")
add_number(doc, "El usuario inicia el rastreo y concede permiso de ubicación en segundo plano.")
add_number(doc, "startLocationUpdatesAsync registra una tarea con distanceInterval de 10 metros.")
add_number(doc, "TaskManager recibe el lote y selecciona la medición más reciente.")
add_number(doc, "expo-device y expo-battery agregan hardware, sistema y nivel de batería.")
add_number(doc, "La aplicación construye el JSON y realiza un POST al webhook de n8n.")

doc.add_heading("Implementación técnica", level=1)
doc.add_heading("Ubicación y marca de tiempo", level=2)
add_body(doc, "La tarea recibe objetos LocationObject. La latitud, longitud, altitud, precisión y velocidad proceden de location.coords. La marca de tiempo se genera a partir de location.timestamp, por lo que conserva el instante de la medición GPS y no solamente el momento en que se ejecutó la petición HTTP.")
doc.add_heading("Hardware y batería", level=2)
add_body(doc, "expo-device obtiene marca, modelo, nombre del sistema y versión. expo-battery devuelve un nivel entre 0 y 1; por ejemplo, 0.85 equivale a 85 %. Cuando el sistema no ofrece un dato opcional, el payload conserva el campo con valor null para mantener un esquema predecible.")
doc.add_heading("Segundo plano", level=2)
add_body(doc, "Android utiliza un servicio en primer plano con notificación persistente. iOS requiere permiso Always, configuración nativa de ubicación en segundo plano y un development build, debido a que Expo Go no admite esta capacidad en iOS. El sistema operativo todavía puede detener el proceso si el usuario fuerza el cierre de la aplicación.")
doc.add_heading("Manejo del webhook", level=2)
add_body(doc, "La respuesta de n8n se procesa inicialmente como texto para aceptar respuestas JSON y cuerpos vacíos. Los estados HTTP no exitosos registran el código y el contenido de respuesta, facilitando el diagnóstico sin confundir un POST exitoso con un error de parseo.")

doc.add_heading("Payload extendido", level=1)
add_code_payload(doc)

doc.add_heading("Descripción de campos", level=2)
fields = doc.add_table(rows=1, cols=3)
set_table_geometry(fields, [2300, 2300, 4760])
headers = ["Bloque", "Campo", "Finalidad"]
for cell, text in zip(fields.rows[0].cells, headers):
    set_cell_shading(cell, "E8EEF5")
    set_font(cell.paragraphs[0].add_run(text), bold=True, color=DARK_BLUE)
field_rows = [
    ("usuario", "usuario", "Identifica al responsable de la medición."),
    ("dispositivo", "nombre / version", "Identifica hardware y sistema operativo."),
    ("ubicacion", "lat / lon", "Representa la posición geográfica."),
    ("ubicacion", "altitud / precision", "Aporta altura y calidad estimada del GPS."),
    ("movimiento", "velocidad", "Permite inferir el tipo de desplazamiento."),
    ("movimiento", "marca_tiempo", "Ordena secuencialmente cada medición."),
    ("estado", "bateria", "Ayuda a evaluar el impacto energético."),
]
for block, field, purpose in field_rows:
    cells = fields.add_row().cells
    for cell, text in zip(cells, (block, field, purpose)):
        set_font(cell.paragraphs[0].add_run(text), color=INK)
set_table_geometry(fields, [2300, 2300, 4760])

doc.add_page_break()
doc.add_heading("Evidencias de funcionamiento", level=1)
add_body(doc, "Sustituye cada recuadro por la captura correspondiente. En Word puedes seleccionar el recuadro, eliminarlo e insertar la imagen desde Insertar > Imágenes. Conserva el texto de figura como descripción.")
add_evidence_box(doc, 1, "Interfaz y permiso de geolocalización", "Aplicación inicializada solicitando o mostrando el permiso de ubicación.")
add_evidence_box(doc, 2, "Historial de ejecuciones de n8n", "Ejecuciones automáticas registradas durante varios tramos de 10 metros.")
add_evidence_box(doc, 3, "JSON extendido recibido", "Payload con altitud, velocidad, timestamp, dispositivo y batería desplegado en n8n.")

doc.add_heading("Procedimiento de prueba", level=1)
add_number(doc, "Abrir la app en un dispositivo físico y habilitar ubicación precisa.")
add_number(doc, "Aceptar el permiso en primer plano y capturar la interfaz.")
add_number(doc, "Presionar Iniciar Rastreo y conceder Permitir siempre.")
add_number(doc, "Caminar más de 30 metros para provocar varias actualizaciones.")
add_number(doc, "Comprobar que el workflow de n8n registra varias ejecuciones.")
add_number(doc, "Expandir un evento y verificar todos los campos del payload.")

doc.add_heading("Resultados esperados", level=1)
add_bullet(doc, "El mapa cambia de posición después del desplazamiento.")
add_bullet(doc, "n8n recibe eventos sin necesidad de polling cada segundo.")
add_bullet(doc, "Cada payload incluye los campos obligatorios de la actividad.")
add_bullet(doc, "La marca de tiempo corresponde a la medición GPS.")
add_bullet(doc, "El botón Detener Rastreo cancela el envío de telemetría.")

doc.add_heading("Conclusiones", level=1)
add_body(doc, "El rastreo basado en distancia reduce transmisiones redundantes frente al polling porque genera telemetría cuando existe movimiento relevante. La integración con Expo permite delegar al sistema operativo la recepción de actualizaciones y mantener funcionamiento en segundo plano bajo las reglas de Android e iOS.")
add_body(doc, "El payload extendido aporta contexto para filtrar mediciones imprecisas, reconstruir una secuencia temporal, distinguir dispositivos y observar el impacto energético. n8n recibe los eventos de forma desacoplada y permite procesarlos posteriormente sin agregar lógica pesada al cliente móvil.")
add_body(doc, "Finalmente, la práctica demuestra que optimizar una aplicación de localización no consiste solamente en reducir peticiones: también requiere permisos transparentes, manejo robusto de errores, trazabilidad temporal y conocimiento de las limitaciones de cada plataforma.")

doc.add_heading("Referencias", level=1)
add_body(doc, "Expo. (2026). Expo Location - SDK 54. https://docs.expo.dev/versions/v54.0.0/sdk/location/")
add_body(doc, "Expo. (2026). Expo Device - SDK 54. https://docs.expo.dev/versions/v54.0.0/sdk/device/")
add_body(doc, "Expo. (2026). Expo Battery - SDK 54. https://docs.expo.dev/versions/v54.0.0/sdk/battery/")
add_body(doc, "Expo. (2026). Expo TaskManager - SDK 54. https://docs.expo.dev/versions/v54.0.0/sdk/task-manager/")

# Evita filas partidas entre páginas y repite el encabezado de la tabla de campos.
for table in doc.tables:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
header_tr_pr = fields.rows[0]._tr.get_or_add_trPr()
tbl_header = OxmlElement("w:tblHeader")
tbl_header.set(qn("w:val"), "true")
header_tr_pr.append(tbl_header)

doc.core_properties.title = "Telemetría de Ubicación por Distancia"
doc.core_properties.subject = "Reporte técnico de MIKELOCATIONS"
doc.core_properties.author = "[ESCRIBE TU NOMBRE COMPLETO]"
doc.core_properties.keywords = "Expo, React Native, ubicación, n8n, telemetría"
doc.save(OUTPUT)
print(OUTPUT)
